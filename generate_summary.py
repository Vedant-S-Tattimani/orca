import os
from docx import Document

def create_summary_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('ORCA: Marine Ecosystem Reasoning with Collaborative Agents', 0)
    title.alignment = 1
    
    doc.add_heading('1. Problem Statement Context (SIH26176 - ISRO)', level=1)
    doc.add_paragraph("The objective of the ORCA platform is to build a system of collaborative AI agents that ingests multi-modal marine data (such as Sea Surface Temperature (SST), chlorophyll levels, ocean winds, and weather forecasts). The system aims to provide intelligent, user-centric, and real-time insights for fishermen, researchers, coastal authorities, and disaster management agencies.")
    
    doc.add_heading('2. Solution Approach', level=1)
    doc.add_paragraph("The solution is architected as a decoupled microservices system. It combines a high-performance backend serving multiple specialized AI agents and a responsive, dynamic frontend.")
    
    doc.add_heading('Backend Architecture', level=2)
    doc.add_paragraph("The backend is built using a multi-agent orchestration pattern:")
    doc.add_paragraph("• Orchestrator Layer: Includes a Planner, Merger, and Location Resolver to dissect user queries, plan which agent(s) to invoke, and merge their responses.", style='List Bullet')
    doc.add_paragraph("• Specialist Agents Layer: Distinct agents for Weather, Sea State, Hazard, and PFZ (Potential Fishing Zones) that independently fetch and process data from real-world APIs like IMD (India Meteorological Department), INCOIS, and ISRO Bhuvan.", style='List Bullet')
    doc.add_paragraph("• Rules Layer (Risk Engine): Implements deterministic safety thresholds (e.g., wind speed caps, wave height limits) to strictly evaluate safety and penalize suitability when dangerous marine conditions are met.", style='List Bullet')
    doc.add_paragraph("• Synthesis Layer: An evidence-based response generator that collates claims from all agents and formulates multi-format response cards.", style='List Bullet')

    doc.add_heading('Frontend Architecture', level=2)
    doc.add_paragraph("The frontend is a fully responsive, lightweight single-page application structure built without heavy frameworks. It utilizes modular JavaScript and TailwindCSS for a modern UI. Key features include:")
    doc.add_paragraph("• dashboard.html: Provides a live operational overview. Connects to real backend APIs to display dynamic risk indices, wind/wave gauges, active directives, and 24-hour weather forecasts.", style='List Bullet')
    doc.add_paragraph("• fishing.html: A dedicated Fishing Intelligence interface with a live interactive map (Leaflet.js). It displays PFZ zones, handles custom point marine safety checks, calculates suitability based on environmental parameters, and provides intelligent 'Safer Area Recommendations' and 'Nearest Port Shelters' during dangerous winds.", style='List Bullet')
    doc.add_paragraph("• assistant.html & map.html: Interfaces for interacting with the AI directly and visualizing extensive tracking routes.", style='List Bullet')
    doc.add_paragraph("• i18n.js, stt.js, tts.js: Integrated multilingual support, Speech-to-Text, and Text-to-Speech capabilities for inclusive accessibility.", style='List Bullet')

    doc.add_heading('3. Technology Stack', level=1)
    doc.add_heading('Backend Stack', level=2)
    doc.add_paragraph("• Framework: Python 3, FastAPI (High-performance Async API)\n• Data Validation: Pydantic\n• Concurrency: AsyncIO\n• HTTP Client: HTTPX (with retry/backoff logic)\n• Geolocation: Geopy / Nominatim", style='List Bullet')
    
    doc.add_heading('Frontend Stack', level=2)
    doc.add_paragraph("• Structure & Styling: HTML5, TailwindCSS (via CDN)\n• Logic: Vanilla JavaScript (ES6+)\n• Mapping & Data Vis: Leaflet.js, Chart.js\n• Typography/Icons: Inter & Geist Mono fonts, Material Symbols", style='List Bullet')

    doc.add_heading('4. Website Deep Dive & Key Features', level=1)
    doc.add_paragraph("The platform ensures maximum utility for the end-user (fishermen/coastal authorities):")
    doc.add_paragraph("• Live Data Integration: No hardcoded fallbacks in production components. The frontend fetches real sensor data via /api/pfz and /api/environmental-data endpoints.", style='List Bullet')
    doc.add_paragraph("• Wind-Penalty Suitability Engine: Even if biological factors (SST/Chlorophyll) are perfect for fishing, the system intelligently caps the fishing suitability to low percentages (e.g., 30%) if wind speeds enter the 'Extreme' or 'High' risk categories (> 40 km/h).", style='List Bullet')
    doc.add_paragraph("• Emergency Recommendations: The frontend analyzes current user coordinates or queried zones against live weather. If dangerous conditions exist, it calculates haversine distances to recommend the nearest safe fishing zone and the closest major port shelter.", style='List Bullet')

    # Save to desktop
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    if not os.path.exists(desktop_path):
        desktop_path = os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop")
        
    filepath = os.path.join(desktop_path, "ORCA_Codebase_Summary.docx")
    doc.save(filepath)
    print(f"Document saved successfully to {filepath}")

if __name__ == '__main__':
    create_summary_doc()
